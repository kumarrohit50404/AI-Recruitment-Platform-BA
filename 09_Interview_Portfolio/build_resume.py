"""Generate modern BA resume Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUT = r"c:\Users\RAHUL\Downloads\Rohit_Kumar_BA_Resume_Modern.docx"
NAVY = "1B2A4A"
GITHUB = "github.com/kumarrohit50404/AI-Recruitment-Platform-BA"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SIDEBAR = RGBColor(0xE8, 0xEC, 0xF0)
SKILL = RGBColor(0xD0, 0xD8, 0xE4)
NAVY_RGB = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2E, 0x6B, 0xE6)
MUTED = RGBColor(0x6B, 0x77, 0x85)
BODY = RGBColor(0x33, 0x33, 0x33)


def shade_cell(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )


def no_borders(table):
    tblPr = table._tbl.tblPr
    tblPr.append(
        parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            "</w:tblBorders>"
        )
    )


def run(p, text, size=10, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def sidebar_title(cell, text, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        p.clear()
    p.paragraph_format.space_before = Pt(14 if not first else 0)
    p.paragraph_format.space_after = Pt(6)
    run(p, text.upper(), 9, True, WHITE)


def sidebar_line(cell, text, size=9):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run(p, text, size, color=SIDEBAR)


def sidebar_skill(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.08)
    run(p, text, 8.5, color=SKILL)


def main_title(cell):
    p = cell.paragraphs[0]
    run(p, "ROHIT KUMAR SRIVASTAVA", 22, True, NAVY_RGB)
    t = cell.add_paragraph()
    run(t, "Junior Business Analyst", 12, color=ACCENT)
    t.paragraph_format.space_after = Pt(10)


def section(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run(p, text.upper(), 11, True, NAVY_RGB)
    line = cell.add_paragraph()
    run(line, "", 4)
    line.paragraph_format.space_after = Pt(6)


def para(cell, text, size=10):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run(p, text, size, color=BODY)


def bullet(cell, text, size=10):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.12)
    for r in p.runs:
        r.text = ""
    run(p, text, size, color=BODY)


def skill_group(cell, title, items):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, 10, True, NAVY_RGB)
    para(cell, items, 9.5)


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)
        s.top_margin = Inches(0.4)
        s.bottom_margin = Inches(0.4)

    table = doc.add_table(1, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(table)
    left, right = table.rows[0].cells[0], table.rows[0].cells[1]
    shade_cell(left, NAVY)

    # --- SIDEBAR ---
    sidebar_title(left, "Contact", first=True)
    for line in [
        "Jamshedpur, Jharkhand",
        "7050139175",
        "kumarrohit50404@gmail.com",
        "Portfolio:",
        GITHUB,
    ]:
        sidebar_line(left, line)

    sidebar_title(left, "Business Analysis")
    for s in [
        "Requirements Elicitation",
        "Stakeholder Analysis",
        "BRD Documentation",
        "User Story Writing",
        "Acceptance Criteria (Gherkin)",
        "Process Modeling (AS-IS / TO-BE)",
        "Requirements Traceability (RTM)",
        "Reporting Requirements",
        "KPI Definition",
        "UAT Planning",
        "Wireframing",
        "MoSCoW Prioritization",
    ]:
        sidebar_skill(left, s)

    sidebar_title(left, "Technical Skills")
    sidebar_line(left, "Excel: Pivot Tables, XLOOKUP,")
    sidebar_line(left, "SUMIFS, Data Validation")
    sidebar_line(left, "SQL: Joins, Aggregations,")
    sidebar_line(left, "KPI Validation Queries")
    sidebar_line(left, "Power BI: KPI Dashboards,")
    sidebar_line(left, "Reporting Requirements")

    sidebar_title(left, "Tools")
    sidebar_line(left, "Confluence | Jira | Figma")
    sidebar_line(left, "Draw.io | Excel | SQL")
    sidebar_line(left, "Power BI | MS Office")

    sidebar_title(left, "Education")
    sidebar_line(left, "MBA, BIT Mesra, Ranchi")
    sidebar_line(left, "2023 - 2025")
    sidebar_line(left, "BBA, NSU Jamshedpur")
    sidebar_line(left, "2018 - 2021")

    sidebar_title(left, "Certifications")
    sidebar_line(left, "SQL for Data Analysis")
    sidebar_line(left, "Agile Fundamentals")
    sidebar_line(left, "Power BI PL-300 (Pursuing)")

    # --- MAIN ---
    main_title(right)

    section(right, "Professional Summary")
    para(
        right,
        "MBA graduate targeting Junior Business Analyst roles. Completed an end-to-end portfolio "
        "case study (HireFlow AI) covering stakeholder analysis, BRD documentation, AS-IS/TO-BE "
        "process modeling, agile user stories, acceptance criteria, requirements traceability, "
        "wireframes, KPI definitions, and executive reporting requirements. Strong analytical "
        "thinking, structured documentation, and stakeholder communication.",
    )

    section(right, "Portfolio Project")
    h = right.add_paragraph()
    run(h, "HireFlow AI - AI-Powered Recruitment Platform", 11, True, NAVY_RGB)
    m = right.add_paragraph()
    run(m, "Business Analyst (Portfolio Case Study) | 2025 - 2026", 9, color=MUTED)
    l = right.add_paragraph()
    run(l, GITHUB, 9, color=ACCENT)
    l.paragraph_format.space_after = Pt(6)

    for b in [
        "Documented business requirements in a structured BRD with MoSCoW prioritization across recruiter, HR, and compliance needs",
        "Conducted stakeholder analysis for eight role groups; mapped pain points to requirements in a stakeholder matrix",
        "Authored user stories and Gherkin acceptance criteria for AI screening, interview scheduling, and reporting workflows",
        "Built requirements traceability matrix linking business needs to user stories and UAT validation scenarios",
        "Modeled AS-IS and TO-BE hiring processes to identify screening and evaluation bottlenecks",
        "Specified wireframes, Figma prototype screens, KPI definitions, and Executive Hiring Dashboard reporting requirements",
        "Validated reporting metrics with sample data: 952 applications, 729 screened, 38 hires, 23.9-day time-to-hire, 80% offer acceptance",
    ]:
        bullet(right, b)

    section(right, "Core Competencies")
    skill_group(
        right,
        "Analysis & Documentation",
        "Requirement Gathering | Gap Analysis | Root Cause Analysis | Business Process Analysis | Stakeholder Communication | Problem Solving",
    )
    skill_group(
        right,
        "Agile & Validation",
        "Agile Fundamentals | Scrum Basics | BRD/FRD Documentation | UAT Support | Decision Support",
    )
    skill_group(
        right,
        "Data & Reporting (BA Scope)",
        "KPI Design | Dashboard Requirements | Data Visualization Concepts | SQL Validation | Excel Reporting",
    )

    section(right, "Internship Experience")
    j = right.add_paragraph()
    run(j, "Vocational Trainee - Tata Steel, Jamshedpur", 10, True, NAVY_RGB)
    d = right.add_paragraph()
    run(d, "May 2024 - June 2024", 9, color=MUTED)
    for b in [
        "Supported data collection, reporting, and operational documentation",
        "Tracked process performance metrics and assisted cross-functional reporting",
        "Contributed to process monitoring and continuous improvement activities",
    ]:
        bullet(right, b)

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    build()
